# -*- coding: utf-8 -*-
"""TSC 기출 1회(유튜브 영상) → 문제·모범답안(3/4/5급) docx. 기존 실전문제집 서식 준용."""
import json, os, sys, importlib.util as _u
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypinyin import pinyin, Style

HERE = os.path.dirname(os.path.abspath(__file__))
TPL = "/Users/leejahun/class-materials/TSC 3급 필수 단어장/TSC 3급 필수 단어장.docx"
GEN = os.path.join(HERE, "gen")
FONT = '맑은 고딕'
NAVY = RGBColor(0x1F,0x3A,0x5F); ORANGE = RGBColor(0xED,0x7D,0x31); DARK = RGBColor(0x33,0x33,0x33)
GRAY = RGBColor(0x7F,0x7F,0x7F); GREENC = RGBColor(0x2E,0x6B,0x2E); BLUE = RGBColor(0x2E,0x74,0xB5)

def py(s):
    return ' '.join(x[0] for x in pinyin(s, style=Style.TONE))

_s = _u.spec_from_file_location("sec", os.path.join(HERE, "sections.py"))
SECM = _u.module_from_spec(_s); _s.loader.exec_module(SECM)
SEC, OPENING, CLOSING = SECM.SEC, SECM.OPENING, SECM.CLOSING
PARTNAME = {"제1부분":"自我介绍 자기소개","제2부분":"看图回答 그림 보고 답하기",
            "제3부분":"快速回答 빠르게 답하기","제4부분":"简短回答 간단히 답하기",
            "제5부분":"拓展回答 논리적으로 답하기","제6부분":"情景应对 상황 대응",
            "제7부분":"看图说话 그림 보고 이야기하기"}

def build(data, ans, out, LV=None):
    d = Document(TPL); body = d.element.body; sect = body.find(qn('w:sectPr'))
    for ch in list(body):
        if ch is not sect: body.remove(ch)
    styles = {s.name for s in d.styles}
    def reloc(el): sect.addprevious(el)
    def kfont(r):
        r.font.name = FONT; pr = r._element.get_or_add_rPr(); rf = pr.find(qn('w:rFonts'))
        if rf is None: rf = OxmlElement('w:rFonts'); pr.append(rf)
        for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rf.set(qn(a), FONT)
    def para(before=0, after=4, align=None):
        p = d.add_paragraph(); reloc(p._element)
        p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
        if align is not None: p.alignment = align
        return p
    def run(p, t, size=11, color=None, bold=False):
        r = p.add_run(t); kfont(r); r.font.size = Pt(size); r.bold = bold
        if color is not None: r.font.color.rgb = color
        return r

    p = para(after=2); run(p, '________ 님    session _____', size=10, color=GRAY)
    _t = f'TSC 기출 1회 — 문제 · 모범답안 ({LV})' if LV else 'TSC 기출 1회 — 문제 · 모범답안 (3급 / 4급 / 5급)'
    p = para(after=1); run(p, _t, size=16, bold=True)
    p = para(after=2); run(p, '출처: TSC 기출문제집 온라인 영상 테스트 (기출 문제 01회)', size=9, color=GRAY)
    p = para(after=12); run(p, ('제1~7부분 전 26문항 + 마무리 발언  ·  질문은 화면·음성 그대로  ·  그림은 원본 상황을 재현한 새 삽화  ·  병음·한국어 병기'
 + (f'  ·  {LV} 답변만 수록' if LV else '')),
                            size=9, color=GRAY)

    p = para(before=8, after=5); run(p, '━━  시험 안내  ━━', size=12, bold=True, color=NAVY)
    for z, k in zip(OPENING["zh"], OPENING["ko"]):
        p = para(after=1); run(p, z, size=10.5, color=DARK)
        p = para(after=4); run(p, '　　', size=9); run(p, k, size=9.5, color=GRAY)

    cur_part = None
    for q in data["questions"]:
        n, part = q["n"], q["part"]
        if part != cur_part:
            cur_part = part
            s = SEC[part]
            p = para(before=16, after=4)
            run(p, f'━━  {part}  {PARTNAME[part]}   ({s["time"]})  ━━', size=12, bold=True, color=NAVY)
            for z, k in zip(s["zh"], s["ko"]):
                p = para(after=1); run(p, z, size=10, color=DARK)
                p = para(after=3); run(p, '　　', size=9); run(p, k, size=9, color=GRAY)
            if s["ex"]:
                qz, a1z, a2z, qk, a1k, a2k = s["ex"]
                p = para(before=3, after=2); run(p, '  채점 기준 예시 (영상 안내에서)', size=9, bold=True, color=ORANGE)
                for z, k, tag in ((qz,qk,''), (a1z,a1k,' ← 짧은 답변'), (a2z,a2k,' ← 자세한 답변, 더 높은 점수')):
                    p = para(after=1); run(p, '  ' + z, size=9.5, color=DARK)
                    if tag: run(p, tag, size=8.5, color=GRAY)
                    p = para(after=3); run(p, '　　　', size=9); run(p, k, size=8.5, color=GRAY)
        p = para(before=10, after=4)
        run(p, f'问题 {n}', size=14, bold=True, color=ORANGE)
        run(p, f'   ({part} 제{q["num"]}문)', size=10, color=GRAY)

        f = os.path.join(GEN, q["img"]["file"]) if q["img"] else None
        if f and os.path.exists(f):
            wid = 5.6 if part == "제7부분" else 2.6
            p = para(after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
            p.add_run().add_picture(f, width=Inches(wid))

        p = para(after=2); run(p, '문제  ', size=9.5, color=NAVY, bold=True)
        run(p, q["zh"] + '  ', size=12, color=DARK, bold=True)
        p = para(after=2); run(p, '　　　', size=9); run(p, py(q["zh"]), size=9, color=GRAY)
        p = para(after=7); run(p, '　　　', size=9); run(p, q["ko"], size=10, color=DARK)

        a = ans[n]
        _lvs = [x for x in (('3급','l3','k3',GREENC), ('4급','l4','k4',BLUE), ('5급','l5','k5',ORANGE))
                if LV is None or x[0] == LV]
        for lv, key, kk, col in _lvs:
            p = para(after=2); run(p, f'모범답안 {lv}  ', size=9.5, color=col, bold=True)
            run(p, a[key], size=11.5, color=DARK, bold=True)
            p = para(after=2); run(p, '　　　', size=9); run(p, py(a[key]), size=8.5, color=GRAY)
            p = para(after=(10 if lv == _lvs[-1][0] else 5)); run(p, '　　　', size=9); run(p, a[kk], size=10, color=DARK)
    p = para(before=16, after=4)
    run(p, '━━  마무리 발언   (발언시간 30초)  ━━', size=12, bold=True, color=NAVY)
    for z, k in zip(CLOSING["zh"], CLOSING["ko"]):
        p = para(after=1); run(p, z, size=10, color=DARK)
        p = para(after=3); run(p, '　　', size=9); run(p, k, size=9, color=GRAY)
    _clv = [x for x in (('3급','l3','k3',GREENC), ('4급','l4','k4',BLUE), ('5급','l5','k5',ORANGE))
            if LV is None or x[0] == LV]
    for lv, key, kk, col in _clv:
        p = para(after=2); run(p, f'모범답안 {lv}  ', size=9.5, color=col, bold=True)
        run(p, CLOSING[key], size=11.5, color=DARK, bold=True)
        p = para(after=2); run(p, '　　　', size=9); run(p, py(CLOSING[key]), size=8.5, color=GRAY)
        p = para(after=(10 if lv==_clv[-1][0] else 5)); run(p, '　　　', size=9); run(p, CLOSING[kk], size=10, color=DARK)
    p = para(before=10, after=2); run(p, CLOSING["end_zh"], size=11, bold=True, color=NAVY)
    p = para(after=4); run(p, CLOSING["end_ko"], size=9.5, color=GRAY)
    d.save(out); print("SAVED:", out)

if __name__ == "__main__":
    data = json.load(open(os.path.join(HERE, "tsc_yt01.json"), encoding="utf-8"))
    s = _u.spec_from_file_location("ans", os.path.join(HERE, "answers.py"))
    m = _u.module_from_spec(s); s.loader.exec_module(m)
    lv = sys.argv[2] if len(sys.argv) > 2 else None
    build(data, m.A, sys.argv[1] if len(sys.argv) > 1 else os.path.join(HERE, "out.docx"), lv)
