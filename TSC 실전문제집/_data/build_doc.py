# -*- coding: utf-8 -*-
# 데이터 기반 TSC 실전테스트 문서 생성기 (JSON + 생성이미지 → docx). 재실행 시 이미지 재활용.
import json,io,sys,os
BASE="/Users/leejahun/class-materials/TSC 실전문제집"
TPL=os.path.join(BASE,"..","TSC 3급 필수 단어장","TSC 3급 필수 단어장.docx")  # LINGO LOUNGE template
DATA=os.path.join(BASE,"_data","tsc_test_data.json")
IMGDIR=os.path.join(BASE,"생성이미지")
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
FONT='맑은 고딕'
NAVY=RGBColor(0x1F,0x3A,0x5F); ORANGE=RGBColor(0xED,0x7D,0x31); DARK=RGBColor(0x33,0x33,0x33)
GRAY=RGBColor(0x7F,0x7F,0x7F); GREENC=RGBColor(0x2E,0x6B,0x2E); RED=RGBColor(0xC0,0x00,0x00)

def build_round(rname, items, outpath):
    d=Document(TPL); body=d.element.body; sect=body.find(qn('w:sectPr'))
    for ch in list(body):
        if ch is not sect: body.remove(ch)
    styles={s.name for s in d.styles}
    def reloc(el): sect.addprevious(el)
    def kfont(r):
        r.font.name=FONT; pr=r._element.get_or_add_rPr(); rf=pr.find(qn('w:rFonts'))
        if rf is None: rf=OxmlElement('w:rFonts'); pr.append(rf)
        for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rf.set(qn(a),FONT)
    def para(style=None,before=0,after=4,align=None):
        p=d.add_paragraph(); reloc(p._element)
        if style and style in styles: p.style=d.styles[style]
        p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
        if align is not None: p.alignment=align
        return p
    def run(p,t,size=11,color=None,bold=False):
        r=p.add_run(t); kfont(r); r.font.size=Pt(size); r.bold=bold
        if color is not None: r.font.color.rgb=color
        return r
    def heading(t):
        p=para('스타일1' if '스타일1' in styles else None,before=12,after=4); run(p,t,size=14,bold=True,color=ORANGE)
    p=para(after=2); run(p,'________ 님    session _____',size=10,color=GRAY)
    p=para(after=1); run(p,f'TSC {rname} — 문제 · 모범답안 정리',size=16,bold=True)
    p=para(after=10); run(p,'제2,3부분 · 문제별 [그림(AI 생성) + 질문(녹음대본) + 모범답안]  ·  그림 병음·한글 병기',size=9,color=GRAY)
    for it in items:
        part=it.get('part','')
        heading(f'问题 {it["n"]}   ({part})' if part else f'问题 {it["n"]}')
        img=os.path.join(IMGDIR,it['image'])
        if os.path.exists(img):
            p=para(after=3,align=WD_ALIGN_PARAGRAPH.CENTER); r=p.add_run(); r.add_picture(img,width=Inches(2.3))
        if it.get('time_label'):
            p=para(after=3,align=WD_ALIGN_PARAGRAPH.CENTER); run(p,'🕐 '+it['time_label'],size=11,color=NAVY,bold=True)
        elif it.get('label'):
            p=para(after=3,align=WD_ALIGN_PARAGRAPH.CENTER); run(p,it['label'],size=11,color=NAVY,bold=True)
        # question
        p=para(after=2); run(p,'질문(녹음대본)  ',size=9.5,color=NAVY,bold=True)
        run(p,it['q_zh']+'  ',size=12,color=DARK,bold=True)
        if it.get('q_py'): run(p,it['q_py'],size=9.5,color=GRAY)
        if it.get('q_ko'):
            p=para(after=6); run(p,'　　　　　　　　',size=9); run(p,it['q_ko'],size=10,color=DARK)
        # answer
        p=para(after=2); run(p,'모범답안  ',size=9.5,color=GREENC,bold=True)
        run(p,it['a_zh']+'  ',size=12,color=DARK,bold=True)
        if it.get('a_py'): run(p,it['a_py'],size=9.5,color=GRAY)
        if it.get('a_ko'):
            p=para(after=(2 if it.get('needs_review') else 12)); run(p,'　　　　　',size=9); run(p,it['a_ko'],size=10,color=DARK)
        if it.get('needs_review'):
            p=para(after=12); run(p,'　　　　　※ 제3부분 녹음대본·모범답안 텍스트는 해설 사진 확인 후 채워 넣을 예정',size=8.5,color=RED)
    d.save(outpath); print("SAVED:",outpath,f"({len(items)} 문항)")

if __name__=="__main__":
    d=json.load(io.open(DATA,encoding='utf-8'))
    rname=sys.argv[1] if len(sys.argv)>1 else "실전1회"
    items=d["rounds"][rname]
    out=os.path.join(BASE,f"TSC {rname}_문제·모범답안.docx")
    build_round(rname, items, out)
