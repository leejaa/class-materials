# -*- coding: utf-8 -*-
"""TSC 실전1~5회 제2~7부분 통합 문서 생성 (마스터/3급/4급).
사용법: python3 build_full.py master   (원본 모범답안)
       python3 build_full.py l3       (3급 쉬운 답변 + 병음)
       python3 build_full.py l4       (4급 답변 + 병음)
제2,3부분=tsc_test_data.json, 제4~7부분=tsc47_content.py. 병음은 pypinyin 자동.
"""
import json, io, sys, os
BASE = "/Users/leejahun/class-materials/TSC 실전문제집"
TPL = os.path.join(BASE, "..", "TSC 3급 필수 단어장", "TSC 3급 필수 단어장.docx")
IMGDIR = os.path.join(BASE, "생성이미지")
sys.path.insert(0, os.path.join(BASE, "_data"))
import tsc47_content as C
import tsc47_levels as L
from pypinyin import pinyin as _py, Style
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT='맑은 고딕'
NAVY=RGBColor(0x1F,0x3A,0x5F); ORANGE=RGBColor(0xED,0x7D,0x31); DARK=RGBColor(0x33,0x33,0x33)
GRAY=RGBColor(0x7F,0x7F,0x7F); GREENC=RGBColor(0x2E,0x6B,0x2E); BLUE=RGBColor(0x1F,0x4E,0x79)

def to_py(zh):
    out=[]
    for seg in _py(zh, style=Style.TONE, errors=lambda x: list(x)):
        out.append(seg[0])
    # 공백 정리: 문장부호 앞 공백 제거
    s=' '.join(out)
    for p in ['，','。','？','！','、','：','；']:
        s=s.replace(' '+p, p).replace(p+' ', p+' ')
    return s

def build(level):
    d=Document(TPL); body=d.element.body; sect=body.find(qn('w:sectPr'))
    for ch in list(body):
        if ch is not sect: body.remove(ch)
    styles={s.name for s in d.styles}
    def reloc(el): sect.addprevious(el)
    def kfont(r):
        r.font.name=FONT; pr=r._element.get_or_add_rPr(); rf=pr.find(qn('w:rFonts'))
        if rf is None: rf=OxmlElement('w:rFonts'); pr.append(rf)
        for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rf.set(qn(a),FONT)
    def para(before=0,after=4,align=None,style=None):
        p=d.add_paragraph(); reloc(p._element)
        if style and style in styles: p.style=d.styles[style]
        p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
        if align is not None: p.alignment=align
        return p
    def run(p,t,size=11,color=None,bold=False,italic=False):
        r=p.add_run(t); kfont(r); r.font.size=Pt(size); r.bold=bold; r.italic=italic
        if color is not None: r.font.color.rgb=color
        return r
    def shade(p,hexcolor):
        pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor); pPr.append(sh)
    def img(name,width=2.3):
        path=os.path.join(IMGDIR,name)
        if os.path.exists(path):
            p=para(after=3,align=WD_ALIGN_PARAGRAPH.CENTER); r=p.add_run(); r.add_picture(path,width=Inches(width))
    def round_head(t):
        p=para(before=16,after=6); shade(p,'1F3A5F'); run(p,'  '+t+'  ',size=17,bold=True,color=RGBColor(0xFF,0xFF,0xFF))
    def part_head(t):
        p=para(before=12,after=4); run(p,'▍'+t,size=13,bold=True,color=ORANGE)
    def q_head(t):
        p=para(before=8,after=3); run(p,t,size=12,bold=True,color=NAVY)
    def ans_block(zh,ko,withpy):
        # 질문/답 공통 중국어+병음+한국어 렌더
        p=para(after=2); run(p,'모범답안  ',size=9.5,color=GREENC,bold=True); run(p,zh,size=12,color=DARK,bold=True)
        if withpy and zh:
            pp=para(after=2); run(pp,to_py(zh),size=9.5,color=GRAY,italic=True)
        if ko:
            pk=para(after=10); run(pk,ko,size=10,color=DARK)
    def q_line(zh,ko,tag='질문(녹음대본)'):
        p=para(after=2); run(p,tag+'  ',size=9.5,color=NAVY,bold=True); run(p,zh,size=12,color=DARK,bold=True)
        if ko:
            pk=para(after=6); run(pk,'　'+ko,size=10,color=GRAY)

    # ---- 데이터 로드 ----
    j=json.load(io.open(os.path.join(BASE,"_data","tsc_test_data.json"),encoding='utf-8'))
    r23=j["rounds"]
    withpy = (level in ('l3','l4'))
    title = {'master':'모범답안(원본)','l3':'3급 쉬운 답변','l4':'4급 답변'}[level]
    def leveled(rnd,pname,n,it):
        if level=='master': return it['a_zh'], it.get('a_ko','')
        e=L.LEVELS.get((rnd,pname,str(n))); fld='l3' if level=='l3' else 'l4'
        if e and e.get(fld): return e[fld], e.get(fld+'_ko','')
        return it['a_zh'], it.get('a_ko','')  # 미작성분은 원본 폴백

    p=para(after=1); run(p,f'TSC 실전테스트 1~5회 · 제2~7부분 — {title}',size=16,bold=True)
    p=para(after=10); run(p,'제2·3부분(그림 단답) → 제4·5부분(구술) → 제6·7부분(상황·스토리)  ·  Lingo Lounge',size=9,color=GRAY)

    rounds=["실전1회","실전2회","실전3회","실전4회","실전5회"]
    for rnd in rounds:
        round_head(f'TSC {rnd}')
        # 제2,3부분
        items=r23.get(rnd,[])
        part_head('제2·3부분 (그림 보고 답하기)')
        for it in items:
            q_head(f'问题 {it["n"]}  ({it.get("part","")})')
            if it.get('image'): img(it['image'])
            if it.get('time_label'):
                pt=para(after=3,align=WD_ALIGN_PARAGRAPH.CENTER); run(pt,'🕐 '+it['time_label'],size=11,color=NAVY,bold=True)
            elif it.get('label'):
                pt=para(after=3,align=WD_ALIGN_PARAGRAPH.CENTER); run(pt,it['label'],size=11,color=NAVY,bold=True)
            q_line(it['q_zh'],it.get('q_ko',''))
            # 제2,3 답은 원본(레벨 무관)
            p=para(after=2); run(p,'모범답안  ',size=9.5,color=GREENC,bold=True); run(p,it['a_zh'],size=12,color=DARK,bold=True)
            pp=para(after=2); run(pp,to_py(it['a_zh']),size=9.5,color=GRAY,italic=True)
            pk=para(after=10); run(pk,it.get('a_ko',''),size=10,color=DARK)
        # 제4~7부분
        cont=C.DATA[rnd]
        for pname in ["제4부분","제5부분","제6부분","제7부분"]:
            block=cont[pname]
            label={'제4부분':'제4부분 (일상 화제 대답)','제5부분':'제5부분 (관점·의견 제시)','제6부분':'제6부분 (상황 대응)','제7부분':'제7부분 (그림 보고 이야기)'}[pname]
            part_head(label)
            if pname=="제7부분":
                if block.get('image'): img(block['image'],width=4.2)
                q_head('问题 ④  (4컷 스토리)')
                pk=para(after=4); run(pk,'　'+block.get('q_ko',''),size=9.5,color=GRAY)
                zh,ko=leveled(rnd,'제7부분','4',block)
                ans_block(zh,ko,withpy)
            else:
                for it in block:
                    tag='질문(상황)' if pname=='제6부분' else '질문(녹음대본)'
                    q_head(f'问题 {it["n"]}')
                    if it.get('image'): img(it['image'])
                    q_line(it['q_zh'],it.get('q_ko',''),tag=tag)
                    zh,ko=leveled(rnd,pname,it['n'],it)
                    ans_block(zh,ko,withpy)

    out=os.path.join(BASE,f'TSC 실전1~5회_제2~7부분_{title}.docx')
    d.save(out); print('SAVED:',out)

if __name__=="__main__":
    build(sys.argv[1] if len(sys.argv)>1 else 'master')
